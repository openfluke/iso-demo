
#!/usr/bin/env python3
"""
telemetry_report_plus.py — Drop‑in upgrade for your telemetry reporting

What’s new vs your current script:
- Loads BOTH public/reports and public/reports_local (and accepts extra paths/argv)
- Adds GPU/CPU speedup, throughput (samples/sec), and per‑model size bucketing (S/M/L/XL)
- Flags anomalies (e.g., “mode collapse”: same prediction for many digits)
- Computes latency percentiles (P50/P90/P99) per machine/model/device
- Produces CSVs alongside the DOCX/PDF for quick spreadsheet work
- Optional PDF export via LibreOffice if present (kept from your script)
- Safer I/O, prettier headings, and more defensive parsing

Usage
-----
python telemetry_report_plus.py
python telemetry_report_plus.py public/reports public/reports_local /some/other/folder

Artifacts
---------
public/manual_reports/report_YYYYMMDD_HHMMSS/
  - telemetry_report.docx
  - telemetry_report.pdf        (if LibreOffice detected)
  - specs.csv
  - summary.csv                 (one row per machine+model)
  - digits.csv                  (per-digit rows)
  - figures/*.png
"""

import os
import re
import sys
import json
import math
import shutil
import statistics as stats
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Tuple, Any, Optional

import numpy as np
import pandas as pd

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import matplotlib
import matplotlib.pyplot as plt

# --------- Config & Paths ---------
PUBLIC_DIR = Path('public')
DEFAULT_INPUT_DIRS = [
    PUBLIC_DIR / 'reports',
    PUBLIC_DIR / 'reports_local',  # also scan local cache by default
]
MANUAL_REPORTS_DIR = PUBLIC_DIR / 'manual_reports'
REPORT_FOLDER = MANUAL_REPORTS_DIR / f'report_{datetime.now().strftime("%Y%m%d_%H%M%S")}'
FIG_DIR = REPORT_FOLDER / 'figures'

# --------- Helpers ---------
def _coerce_float(x, default=np.nan):
    try:
        return float(x)
    except Exception:
        return default

def _model_size_bucket(name: str) -> str:
    # Heuristic: mnist_S1.json → S, mnist_XL2.json → XL
    m = re.search(r'_(S|M|L|XL)\d*\.json$', name, re.IGNORECASE)
    if not m:
        m = re.search(r'(S|M|L|XL)', name, re.IGNORECASE)
    return (m.group(1).upper() if m else 'UNK')

def _safe_get(d: dict, *keys, default=None):
    cur = d
    for k in keys:
        if cur is None or (not isinstance(cur, dict)) or k not in cur:
            return default
        cur = cur[k]
    return cur

def _pct(a: Iterable[float], q: float) -> float:
    arr = [x for x in a if isinstance(x, (int, float)) and not math.isnan(x)]
    if not arr:
        return np.nan
    arr = sorted(arr)
    idx = int(round((q/100.0)*(len(arr)-1)))
    return float(arr[idx])

def _is_mode_collapse(per_digit_preds: List[int]) -> bool:
    # If 7+ out of 10 digits map to the same predicted label, flag as suspicious
    if not per_digit_preds:
        return False
    most = max((per_digit_preds.count(v) for v in set(per_digit_preds)))
    return most >= 7

# --------- Load Telemetry ---------
def find_input_dirs(argv: List[str]) -> List[Path]:
    if argv:
        return [Path(a) for a in argv]
    return DEFAULT_INPUT_DIRS

def discover_json_files(input_dirs: List[Path]) -> List[Path]:
    files: List[Path] = []
    for d in input_dirs:
        if d.exists():
            files.extend(sorted(d.glob('telemetry_*.json')))
    # Unique by stem
    seen = set()
    uniq = []
    for f in files:
        if f.stem not in seen:
            uniq.append(f)
            seen.add(f.stem)
    return uniq

def load_all() -> Tuple[pd.DataFrame, pd.DataFrame, List[Path]]:
    json_files = discover_json_files(find_input_dirs(sys.argv[1:]))
    if not json_files:
        raise SystemExit("No telemetry_*.json files found in default or provided folders.")

    specs_rows: List[Dict[str, Any]] = []
    perf_rows: List[Dict[str, Any]] = []

    for fp in json_files:
        try:
            with open(fp, 'r') as f:
                t = json.load(f)
        except Exception as e:
            print(f"⚠️  Skipping unreadable {fp}: {e}")
            continue

        sysinfo = t.get('system_info', {})
        machine_id = t.get('machine_id', fp.stem)
        started_at = t.get('started_at', '')
        machine_name = sysinfo.get('device_model') or sysinfo.get('cpu_model', 'unknown').split(' CPU')[0]

        specs_rows.append({
            'machine_id': machine_id,
            'machine_name': machine_name,
            'architecture': sysinfo.get('architecture', ''),
            'os': sysinfo.get('os', ''),
            'os_version': sysinfo.get('os_version', ''),
            'cpu_model': sysinfo.get('cpu_model', ''),
            'gpu_model': sysinfo.get('gpu_model', ''),
            'ram_gb': float(_safe_get(sysinfo, 'ram_bytes', default=0)) / (1024**3),
            'started_at': started_at,
            'models_tested': len(t.get('per_model', [])),
        })

        for model in t.get('per_model', []):
            model_file = model.get('model_file', '')
            size_bucket = _model_size_bucket(model_file)
            webgpu_ok = bool(model.get('webgpu_init_ok', False))
            webgpu_init_time_ms = _coerce_float(model.get('webgpu_init_time_ms'))

            cpu_recs = model.get('cpu', []) or []
            gpu_recs = model.get('gpu', []) or []
            drift_recs = model.get('drift', []) or []

            # Per-model summary (10 digits)
            cpu_lat = [_coerce_float(r.get('elapsed_ms')) for r in cpu_recs]
            gpu_lat = [_coerce_float(r.get('elapsed_ms')) for r in gpu_recs] if webgpu_ok else []
            adhd10 = model.get('adhd10', {}) or {}

            # Speedup & throughput
            cpu_ms = np.nanmean(cpu_lat) if cpu_lat else np.nan
            gpu_ms = np.nanmean(gpu_lat) if gpu_lat else np.nan
            speedup = (cpu_ms / gpu_ms) if (isinstance(cpu_ms, float) and isinstance(gpu_ms, float) and gpu_ms and not math.isnan(cpu_ms) and not math.isnan(gpu_ms)) else np.nan
            cpu_throughput = (1000.0 / cpu_ms) if cpu_ms and not math.isnan(cpu_ms) else np.nan
            gpu_throughput = (1000.0 / gpu_ms) if gpu_ms and not math.isnan(gpu_ms) else np.nan

            perf_rows.append({
                'machine_id': machine_id,
                'machine_name': machine_name,
                'model_file': model_file,
                'model_size': size_bucket,
                'webgpu_init_ok': webgpu_ok,
                'webgpu_init_time_ms': webgpu_init_time_ms,
                'cpu_top1_accuracy': _coerce_float(adhd10.get('top1_accuracy_cpu')),
                'gpu_top1_accuracy': _coerce_float(adhd10.get('top1_accuracy_gpu')),
                'cpu_vs_gpu_agree_count': _coerce_float(adhd10.get('cpu_vs_gpu_agree_count')),
                'avg_drift_mae': _coerce_float(adhd10.get('avg_drift_mae')),
                'max_drift_max_abs': _coerce_float(adhd10.get('max_drift_max_abs')),
                'cpu_elapsed_avg_ms': cpu_ms,
                'gpu_elapsed_avg_ms': gpu_ms,
                'speedup_gpu_over_cpu': speedup,
                'cpu_samples_per_sec': cpu_throughput,
                'gpu_samples_per_sec': gpu_throughput,
                'lat_p50_cpu_ms': _pct(cpu_lat, 50),
                'lat_p90_cpu_ms': _pct(cpu_lat, 90),
                'lat_p99_cpu_ms': _pct(cpu_lat, 99),
                'lat_p50_gpu_ms': _pct(gpu_lat, 50) if webgpu_ok else np.nan,
                'lat_p90_gpu_ms': _pct(gpu_lat, 90) if webgpu_ok else np.nan,
                'lat_p99_gpu_ms': _pct(gpu_lat, 99) if webgpu_ok else np.nan,
                'is_digit': False,
            })

            # Per-digit rows
            for i, cpu_d in enumerate(cpu_recs):
                gpu_d = (gpu_recs[i] if i < len(gpu_recs) else {}) if webgpu_ok else {}
                drift_d = (drift_recs[i] if i < len(drift_recs) else {}) or {}
                perf_rows.append({
                    'machine_id': machine_id,
                    'machine_name': machine_name,
                    'model_file': model_file,
                    'model_size': size_bucket,
                    'is_digit': True,
                    'digit': cpu_d.get('digit'),
                    'idx': cpu_d.get('idx'),
                    'cpu_pred': cpu_d.get('pred'),
                    'gpu_pred': gpu_d.get('pred', np.nan),
                    'cpu_top1_score': _coerce_float(cpu_d.get('top1_score')),
                    'gpu_top1_score': _coerce_float(gpu_d.get('top1_score', np.nan)),
                    'cpu_elapsed_ms': _coerce_float(cpu_d.get('elapsed_ms')),
                    'gpu_elapsed_ms': _coerce_float(gpu_d.get('elapsed_ms', np.nan)),
                    'drift_mae': _coerce_float(drift_d.get('mae')),
                    'drift_max_abs': _coerce_float(drift_d.get('max_abs')),
                })

    df_specs = pd.DataFrame(specs_rows)
    df_perf = pd.DataFrame(perf_rows)

    # Anomaly flags: per machine+model, check if many digits predict same class
    if not df_perf.empty:
        collapses = []
        for (mach, mod), sub in df_perf[df_perf['is_digit'] == True].groupby(['machine_name', 'model_file']):
            preds = [int(p) for p in sub['cpu_pred'].dropna().tolist() if isinstance(p, (int, np.integer))]
            if _is_mode_collapse(preds):
                collapses.append({'machine_name': mach, 'model_file': mod, 'anomaly': 'cpu_mode_collapse'})
            if 'gpu_pred' in sub.columns:
                gp = [int(p) for p in sub['gpu_pred'].dropna().tolist() if isinstance(p, (int, np.integer))]
                if _is_mode_collapse(gp):
                    collapses.append({'machine_name': mach, 'model_file': mod, 'anomaly': 'gpu_mode_collapse'})
        df_anom = pd.DataFrame(collapses)
    else:
        df_anom = pd.DataFrame(columns=['machine_name', 'model_file', 'anomaly'])

    return df_specs, df_perf, json_files, df_anom

# --------- Reporting ---------
def add_specs_table(doc: Document, df_specs: pd.DataFrame):
    doc.add_heading('Machine Specifications', level=1)
    cols = ['machine_name', 'architecture', 'os', 'os_version', 'cpu_model', 'gpu_model', 'ram_gb', 'models_tested', 'started_at']
    df = df_specs[cols].rename(columns={
        'machine_name': 'Machine',
        'architecture': 'Arch',
        'os': 'OS',
        'os_version': 'OS Version',
        'cpu_model': 'CPU',
        'gpu_model': 'GPU',
        'ram_gb': 'RAM (GB)',
        'models_tested': 'Models',
        'started_at': 'Started',
    })
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, c in enumerate(df.columns):
        hdr = table.rows[0].cells[i]
        hdr.text = c
        hdr.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(round(v, 2) if isinstance(v, (float, int)) else v)
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

def savefig(doc: Document, fig, title: str, fname: str):
    out = FIG_DIR / f'{fname}.png'
    fig.savefig(out, dpi=300, bbox_inches='tight')
    plt.close(fig)
    doc.add_heading(title, level=2)
    doc.add_picture(str(out), width=Inches(6))

def add_performance_figures(doc: Document, df_perf: pd.DataFrame):
    doc.add_heading('Performance Comparisons', level=1)
    df_sum = df_perf[df_perf['is_digit'] == False].copy()
    if df_sum.empty:
        doc.add_paragraph('No summary rows parsed.')
        return

    # Average latency bars
    for fld, title, fname in [
        ('cpu_elapsed_avg_ms', 'Average CPU Inference Time per Model (ms)', 'avg_cpu_time'),
        ('gpu_elapsed_avg_ms', 'Average GPU Inference Time per Model (ms)', 'avg_gpu_time'),
        ('cpu_top1_accuracy', 'Top‑1 Accuracy (CPU) per Model', 'acc_cpu'),
        ('gpu_top1_accuracy', 'Top‑1 Accuracy (GPU) per Model', 'acc_gpu'),
        ('avg_drift_mae', 'CPU↔GPU Drift (MAE) per Model', 'drift_mae'),
        ('cpu_vs_gpu_agree_count', 'CPU‑GPU Agreement Count (out of 10)', 'agreement'),
        ('webgpu_init_time_ms', 'WebGPU Init Time per Machine (ms)', 'webgpu_init'),
        ('speedup_gpu_over_cpu', 'GPU/CPU Speedup (×)', 'speedup'),
        ('gpu_samples_per_sec', 'GPU Throughput (samples/sec)', 'throughput_gpu'),
    ]:
        sub = df_sum.copy()
        sub = sub.replace([np.inf, -np.inf], np.nan)
        # Be explicit to avoid the FutureWarning about downcasting
        pd.set_option('future.no_silent_downcasting', True)
        sub = sub.dropna(subset=['machine_name', 'model_file', fld], how='any')
        if sub.empty:
            continue
        fig, ax = plt.subplots(figsize=(10, 6))
        # Minimal matplotlib (no seaborn) for portability
        # Use a pivot_table to aggregate duplicates robustly and avoid reindex errors
        piv = sub.pivot_table(index='machine_name', columns='model_file', values=fld, aggfunc='mean')
        # Freeze orderings for stable chart layout
        machines = list(piv.index)
        models = list(piv.columns)
        width = 0.8 / max(1, len(models))
        x = np.arange(len(machines))
        for j, m in enumerate(models):
            y = piv[m].astype(float).values
            ax.bar(x + j*width, y, width=width, label=m)
        ax.set_xticks(x + (len(models)-1)*width/2)
        ax.set_xticklabels(machines, rotation=30, ha='right')
        ax.set_title(title)
        ax.legend(fontsize=8, ncols=2)
        savefig(doc, fig, title, fname)

    # Percentile latency scatter (CPU vs GPU)
    sub = df_sum.replace([np.inf, -np.inf], np.nan)
    if {'lat_p90_cpu_ms','lat_p90_gpu_ms'}.issubset(sub.columns):
        gg = sub.dropna(subset=['lat_p90_cpu_ms','lat_p90_gpu_ms'])
        if not gg.empty:
            fig, ax = plt.subplots(figsize=(7, 6))
            ax.scatter(gg['lat_p90_cpu_ms'], gg['lat_p90_gpu_ms'])
            # After (cleaner)
            for _, r in gg.iterrows():
                label = f"{str(r['machine_id'])[:6]} ({r['model_size']})"
                # Only annotate points in the top 5% of GPU P90 or CPU P90
                if (r['lat_p90_cpu_ms'] > gg['lat_p90_cpu_ms'].quantile(0.95) or
                    r['lat_p90_gpu_ms'] > gg['lat_p90_gpu_ms'].quantile(0.95)):
                    ax.annotate(label,
                                (r['lat_p90_cpu_ms'], r['lat_p90_gpu_ms']),
                                fontsize=7, alpha=0.7)

            ax.set_xlabel('CPU P90 (ms)')
            ax.set_ylabel('GPU P90 (ms)')
            ax.set_title('Latency P90: CPU vs GPU')
            savefig(doc, fig, 'Latency P90: CPU vs GPU', 'p90_scatter')

    # Per‑digit heatmaps (CPU time and drift)
    df_digit = df_perf[df_perf['is_digit'] == True].copy()
    if not df_digit.empty:
        # CPU time heatmap
        piv = df_digit.pivot_table(values='cpu_elapsed_ms', index='machine_name', columns='digit', aggfunc='mean')
        if not piv.empty:
            fig, ax = plt.subplots(figsize=(12, 7))
            im = ax.imshow(piv.values, aspect='auto')
            ax.set_xticks(range(len(piv.columns)))
            ax.set_xticklabels(piv.columns)
            ax.set_yticks(range(len(piv.index)))
            ax.set_yticklabels(piv.index)
            ax.set_title('CPU Elapsed Time (ms) per Digit')
            fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
            savefig(doc, fig, 'CPU Elapsed Time Heatmap per Digit', 'heat_cpu_digit')

        # Drift heatmap
        piv2 = df_digit.dropna(subset=['drift_max_abs']).pivot_table(values='drift_max_abs', index='machine_name', columns='digit', aggfunc='mean')
        if not piv2.empty:
            fig, ax = plt.subplots(figsize=(12, 7))
            im = ax.imshow(piv2.values, aspect='auto')
            ax.set_xticks(range(len(piv2.columns)))
            ax.set_xticklabels(piv2.columns)
            ax.set_yticks(range(len(piv2.index)))
            ax.set_yticklabels(piv2.index)
            ax.set_title('Max Abs Drift per Digit')
            fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
            savefig(doc, fig, 'Max Abs Drift Heatmap per Digit', 'heat_drift_digit')

def add_summary_table(doc: Document, df_sum: pd.DataFrame, df_anom: pd.DataFrame):
    doc.add_heading('Summary Performance Table', level=2)
    cols = [
        'machine_name','model_file','model_size','cpu_top1_accuracy','gpu_top1_accuracy',
        'avg_drift_mae','webgpu_init_time_ms','speedup_gpu_over_cpu'
    ]
    sub = df_sum[cols].copy()
    sub = sub.replace([np.inf, -np.inf], np.nan)

    table = doc.add_table(rows=1, cols=len(cols))
    table.style = 'Table Grid'
    for i, h in enumerate(['Machine','Model','Size','CPU Acc','GPU Acc','Drift MAE','Init (ms)','Speedup×']):
        cell = table.rows[0].cells[i]
        cell.text = h

    for _, r in sub.iterrows():
        cells = table.add_row().cells
        vals = [
            r['machine_name'], r['model_file'], r['model_size'],
            f"{r['cpu_top1_accuracy']:.3f}" if pd.notna(r['cpu_top1_accuracy']) else '–',
            f"{r['gpu_top1_accuracy']:.3f}" if pd.notna(r['gpu_top1_accuracy']) else '–',
            f"{r['avg_drift_mae']:.6f}" if pd.notna(r['avg_drift_mae']) else '–',
            f"{r['webgpu_init_time_ms']:.1f}" if pd.notna(r['webgpu_init_time_ms']) else '–',
            f"{r['speedup_gpu_over_cpu']:.2f}" if pd.notna(r['speedup_gpu_over_cpu']) else '–',
        ]
        for i, v in enumerate(vals):
            cells[i].text = str(v)

    if not df_anom.empty:
        doc.add_paragraph()
        doc.add_paragraph("⚠ Anomalies detected:")
        for _, row in df_anom.iterrows():
            doc.add_paragraph(f"• {row['machine_name']} / {row['model_file']}: {row['anomaly']}", style=None)

def write_sidecar_csvs(df_specs: pd.DataFrame, df_perf: pd.DataFrame):
    (REPORT_FOLDER).mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    df_specs.to_csv(REPORT_FOLDER / 'specs.csv', index=False)
    df_perf.to_csv(REPORT_FOLDER / 'all_rows.csv', index=False)
    df_perf[df_perf['is_digit']==False].to_csv(REPORT_FOLDER / 'summary.csv', index=False)
    df_perf[df_perf['is_digit']==True].to_csv(REPORT_FOLDER / 'digits.csv', index=False)

def maybe_export_pdf(docx_path: Path):
    # Try LibreOffice headless; if not found, skip gracefully.
    try:
        import shutil
        lo = shutil.which('libreoffice') or shutil.which('soffice')
        if lo:
            os.system(f'"{lo}" --headless --convert-to pdf --outdir "{docx_path.parent}" "{docx_path}"')
            print(f"PDF generated: {docx_path.with_suffix('.pdf')}")
        else:
            print("LibreOffice not found — skipping PDF export. (Install libreoffice to enable.)")
    except Exception as e:
        print(f"PDF export failed: {e}")

def main():
    MANUAL_REPORTS_DIR.mkdir(exist_ok=True)
    REPORT_FOLDER.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)

    df_specs, df_perf, json_files, df_anom = load_all()
    write_sidecar_csvs(df_specs, df_perf)

    # Build the document
    doc = Document()
    doc.add_heading('Distributed ML Infrastructure Testing Framework — Telemetry Report', 0)
    doc.add_paragraph(f'Generated on {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'Analyzed {len(json_files)} telemetry files from {df_specs["machine_id"].nunique()} machine(s).')
    doc.add_paragraph(f'Total model evaluations (summary rows): {len(df_perf[df_perf["is_digit"]==False])}')
    doc.add_paragraph()

    add_specs_table(doc, df_specs)
    add_performance_figures(doc, df_perf)
    add_summary_table(doc, df_perf[df_perf["is_digit"]==False].copy(), df_anom)

    docx_out = REPORT_FOLDER / 'telemetry_report.docx'
    doc.save(docx_out)
    print(f"DOCX generated: {docx_out}")

    # Optional PDF
    maybe_export_pdf(docx_out)

if __name__ == "__main__":
    main()